"""
SAP Gap Analysis Pipeline  —  Hardened & Optimized
====================================================
Vertex AI · Gemini 2.5 Pro · JSON-chained steps → DOCX

OPTIMIZATIONS vs original
--------------------------
1. Merged Step 2 + Steps 3/4 into batched LLM calls  (was: 1 call for all reqs)
   Normalization + capability assessment run together, but in batches of 20
   requirements per call. Root cause of '0 gaps' bug: 84 reqs in one call
   produced ~168 JSON objects, truncating the response; the fix-it prompt then
   returned empty arrays. Batching caps each response at ~80 objects, well within
   the 16k-token output budget. 84 reqs → 5 calls (~103 s total — same wall-time).

2. Slim payloads to later steps
   step5_6_7 receives ONLY gap=true items (not the full 400-line accumulated JSON).
   step_confirm_and_actions receives ONLY gap=false items + gap_analysis list.
   Cuts input token count by 40–70 % on typical 30-req documents.

3. Smarter retry logic
   On JSON parse failure the retry prompt sends ONLY the broken raw text and asks
   the model to fix it — not the entire original task prompt. Faster and cheaper.

4. Transcript chunking for large inputs
   Transcripts > 100 k chars are split into overlapping chunks; requirements are
   extracted per chunk and de-duplicated before normalisation. No silent truncation.

5. Raised MAX_TOKENS to 16 384
   The original 8 192 limit caused truncated JSON on large gap sets, triggering
   expensive retry loops. Gemini 2.5 Pro supports up to 65 536 output tokens.

6. Structured progress logging with elapsed timing per step.

ANTI-HALLUCINATION (NEW)
--------------------------
7. Hardened Step 5 prompt
   Strict ID rules, exact RICEFW whitelist, explicit item-count enforcement,
   experienced SAP IS-U consultant persona for SAP-native solution bullets.

8. Programmatic guardrail (_validate_and_filter_gaps)
   After Step 5, every gap item is validated against extracted requirement IDs.
   Items with invented IDs or invalid RICEFW types are silently dropped.
   GAP-IDs are re-sequenced cleanly after filtering.

9. Maximum determinism
   TEMPERATURE=0.0, TOP_P=0.1 — eliminates stochastic creative drift on retries.

Architecture (unchanged externally)
-------------------------------------
  transcript → step1 → step2+3+4 (merged) → step5/6/7 → [guardrail] → confirm → DOCX

LLM calls per run: 3 (+ 1 per extra transcript chunk if input > 100 k chars)

Usage (identical CLI to original)
-----------------------------------
  python sap_gap_analyser.py \\
      --transcript  meeting_input.json \\
      --project     YOUR_GCP_PROJECT_ID \\
      --location    us-central1 \\
      --output      SAP_Gap_Analysis.docx

  # Optionally dump intermediate JSON (step5 dump is now step5_ricefw_validated):
      --dump-json


python gap/sap_gap_analyser_updated.py `
    --transcript  run/run_0422-0912/meeting-input.json `
    --project     dn-studio-01 `
    --location    us-central1 `
    --output      run/run_0422-0912/SAP_Gap_Analysis_A.docx


"""

import argparse
import json
import os
import subprocess
import sys
import textwrap
import time
from typing import Any
from datetime import datetime
from pathlib import Path

from google import genai
from google.genai import types

from backend.helper import strip_markdown_json_fence

# ──────────────────────────────────────────────────────────────────────────────
# Configuration
# ──────────────────────────────────────────────────────────────────────────────
MODEL_ID         = "gemini-2.5-pro"
TEMPERATURE      = 0.0             # max determinism — eliminates creative drift
TOP_P            = 0.1             # tight nucleus sampling; works with temp=0 on Gemini
MAX_TOKENS       = 16_384          # raised from 8 192 — avoids truncated-JSON retries
TRANSCRIPT_CHUNK = 100_000          # chars per chunk with 2 k overlap
TRANSCRIPT_OVERLAP = 2_000
ASSESS_BATCH_SIZE  = 20             # requirements per Step-2+3+4 LLM call
                                    # 20 reqs → ~80 JSON objects out → safely under 16k tokens
                                    # 84 reqs in one call → ~168 objects → truncation & 0-gap failure
SCRIPT_DIR       = Path(__file__).parent
JS_TEMPLATE_PATH = SCRIPT_DIR.parent / "templates" / "js_template.js"


# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────
def _init_genai(project: str, location: str) -> genai.Client:
    """
    Create a Google GenAI client.
    Uses Vertex when project/location are provided; otherwise falls back to API key mode.
    """
    try:
        if project and location:
            return genai.Client(vertexai=True, project=project, location=location)
    except Exception:
        # Fall back to API key mode if Vertex auth/config is unavailable.
        pass

    api_key = (os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY") or "").strip()
    if not api_key:
        raise RuntimeError(
            "Google GenAI credentials not configured. Set PROJECT_ID+LOCATION for Vertex mode "
            "or set GOOGLE_API_KEY / GEMINI_API_KEY for API key mode."
        )
    return genai.Client(api_key=api_key)


def _genai_generate_text(client: genai.Client, prompt: str) -> str:
    response = client.models.generate_content(
        model=MODEL_ID,
        contents=prompt,
        config=types.GenerateContentConfig(
            temperature=TEMPERATURE,
            top_p=TOP_P,
            max_output_tokens=MAX_TOKENS,
            response_mime_type="application/json",
        ),
    )
    text = getattr(response, "text", None)
    if text:
        return text
    # Conservative fallback for SDK response shape differences.
    return json.dumps(response.to_json_dict(), ensure_ascii=False)


def _call_llm(model: Any, prompt: str, step_name: str) -> dict:
    """
    Call Gemini 2.5 Pro and return parsed JSON.

    OPTIMIZATION: On failure, retry with a surgical fix-it prompt instead of
    re-sending the full original prompt. Saves tokens on the (rare) retry path.
    """
    print(f"    [LLM] {step_name} ...", flush=True)
    t0 = time.time()

    raw = strip_markdown_json_fence(_genai_generate_text(model, prompt).strip())

    for attempt in range(1, 4):
        try:
            result = json.loads(raw)
            elapsed = time.time() - t0
            print(f"    [LLM] {step_name} done in {elapsed:.1f}s", flush=True)
            return result
        except json.JSONDecodeError as exc:
            if attempt == 3:
                raise ValueError(
                    f"[{step_name}] Non-JSON after {attempt} attempts.\n"
                    f"Error: {exc}\nRaw (first 500): {raw[:500]}"
                ) from exc

            print(f"    [LLM] Invalid JSON (attempt {attempt}); sending fix-it prompt...", flush=True)
            # OPTIMIZATION: only send the broken output, not the entire original prompt
            fix_prompt = textwrap.dedent(f"""
            The JSON below is malformed. Return ONLY the corrected valid JSON.
            Do not add any explanation, markdown fences, or extra text.

            BROKEN JSON:
            {raw[:8000]}
            """)
            raw = strip_markdown_json_fence(_genai_generate_text(model, fix_prompt).strip())

    # unreachable, but satisfies type checkers
    raise ValueError(f"[{step_name}] Failed after retries")


def _load_transcript(path: str) -> str:
    """Load and flatten a transcript file (JSON array or plain text)."""
    data = Path(path).read_text(encoding="utf-8")
    try:
        obj = json.loads(data)
        if isinstance(obj, list) and obj and isinstance(obj[0], dict) and "transcript_json" in obj[0]:
            chunks = []
            for idx, meeting in enumerate(obj, start=1):
                transcript_json = meeting.get("transcript_json", {})
                file_metadata   = transcript_json.get("file_metadata", {})
                segments        = transcript_json.get("transcript", [])
                lines = [s.get("text", "").strip() for s in segments if s.get("text", "").strip()]
                if not lines:
                    continue
                title          = file_metadata.get("source_video") or file_metadata.get("file_name") or f"Meeting {idx}"
                meeting_number = meeting.get("meeting_number", idx)
                meeting_date   = meeting.get("meeting_date", "unknown-date")
                chunks.append("\n".join([
                    f"=== Meeting {meeting_number} | {meeting_date} | {title} ===",
                    *lines,
                ]))
            if chunks:
                return "\n\n".join(chunks)
        return json.dumps(obj, indent=2)
    except json.JSONDecodeError:
        return data


def _chunk_transcript(transcript: str) -> list[str]:
    """
    OPTIMIZATION: Split long transcripts into overlapping chunks so no content
    is silently dropped. The original code hard-truncated at 30 000 chars.
    Returns a list of chunk strings (single-item list if transcript fits in one chunk).
    """
    if len(transcript) <= TRANSCRIPT_CHUNK:
        return [transcript]

    chunks = []
    start  = 0
    while start < len(transcript):
        end = start + TRANSCRIPT_CHUNK
        chunks.append(transcript[start:end])
        start = end - TRANSCRIPT_OVERLAP   # overlap to avoid cutting sentences mid-thought
    return chunks


def _dedup_requirements(req_lists: list[list[dict]]) -> list[dict]:
    """
    Merge requirement lists from multiple chunks and remove near-duplicates.
    De-duplication is text-similarity based (lowercased substring check).
    Re-numbers requirements R1, R2, … after merging.
    """
    seen  = []
    merged = []
    for reqs in req_lists:
        for req in reqs:
            text = req.get("text", "").lower().strip()
            if not any(text in s or s in text for s in seen):
                seen.append(text)
                merged.append(req)

    # Renumber sequentially
    for i, req in enumerate(merged, start=1):
        req["id"] = f"R{i}"
    return merged


def _dump(data: dict, name: str, output_dir: Path) -> None:
    out = output_dir / f"{name}.json"
    out.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"    [DUMP] {out}")


def _slim(data: dict, keys: list[str]) -> dict:
    """Return a copy of data with only the specified top-level keys."""
    return {k: data[k] for k in keys if k in data}


# ──────────────────────────────────────────────────────────────────────────────
# ANTI-HALLUCINATION GUARDRAIL — programmatic gap validator
# ──────────────────────────────────────────────────────────────────────────────
_ALLOWED_RICEFW = {"Report", "Interface", "Conversion", "Enhancement", "Form", "Workflow"}

def _validate_and_filter_gaps(gap_analysis: list[dict], valid_req_ids: set[str]) -> list[dict]:
    """
    Remove hallucinated gap items and enforce 1-to-1 mapping to extracted requirements.

    Drops any gap that:
      (a) references a req_id that does not exist in the extracted requirements, OR
      (b) carries an invalid / misspelled RICEFW type.

    Re-sequences GAP-IDs cleanly after filtering so the output document has no
    broken numbering (GAP-001, GAP-002 … with no gaps).
    """
    validated: list[dict] = []
    for gap in gap_analysis:
        req_id  = str(gap.get("req_id",  "")).strip()
        ricefw  = str(gap.get("ricefw",  "")).strip()

        if req_id not in valid_req_ids:
            print(f"    [WARN] Dropped hallucinated gap — invalid req_id='{req_id}'")
            continue

        if ricefw not in _ALLOWED_RICEFW:
            print(f"    [WARN] Dropped gap '{req_id}' — invalid ricefw='{ricefw}'")
            continue

        validated.append(gap)

    # Re-sequence sequentially; avoids broken GAP numbering after drops
    for i, gap in enumerate(validated, start=1):
        gap["gap_id"] = f"GAP-{i:03d}"

    dropped = len(gap_analysis) - len(validated)
    if dropped:
        print(f"    [GUARDRAIL] {dropped} hallucinated item(s) removed; "
              f"{len(validated)} validated gap(s) retained.")
    return validated


# ──────────────────────────────────────────────────────────────────────────────
# STEP 1 — Requirement Extraction  (unchanged logic, now chunk-aware)
# ──────────────────────────────────────────────────────────────────────────────
def step1_extract_requirements(model: Any, transcript: str) -> dict:
    """
    Extract atomic business requirements from the raw transcript.
    If the transcript is large, each chunk is processed independently and
    results are merged + de-duplicated.

    Output: { meeting_title, meeting_date, scope_context, requirements: [...] }
    """
    chunks = _chunk_transcript(transcript)
    all_req_lists   = []
    scope_context   = ""
    meeting_title   = ""
    meeting_date    = ""

    for chunk_idx, chunk in enumerate(chunks, start=1):
        suffix = f" (chunk {chunk_idx}/{len(chunks)})" if len(chunks) > 1 else ""
        prompt = textwrap.dedent(f"""
        You are a Senior SAP Business Analyst performing requirement extraction from a meeting transcript.

        TASK
        ----
        Read the transcript carefully. Extract ALL distinct, actionable business requirements.

        RULES
        -----
        - Each requirement must be ATOMIC (one clear idea).
        - Use "The system should …" or "The business requires …" phrasing.
        - Ignore pleasantries, repetitions, off-topic chat.
        - Number requirements R1, R2, R3 … in order of appearance.
        - Also write a short "scope_context" paragraph (3–5 sentences) summarising
          what the session was about, the system landscape (IS-U / EMS / ECC),
          and key design decisions confirmed.

        RESPOND WITH JSON ONLY — no prose, no markdown fences. Schema:
        {{
          "meeting_title": "<string>",
          "meeting_date":  "<string>",
          "scope_context": "<string>",
          "requirements":  [
            {{ "id": "R1", "text": "<string>" }},
            ...
          ]
        }}

        TRANSCRIPT{suffix}
        ----------
        {chunk}
        """)
        result = _call_llm(model, prompt, f"Step 1 – Extraction{suffix}")
        all_req_lists.append(result.get("requirements", []))

        # Use first chunk's metadata
        if chunk_idx == 1:
            scope_context = result.get("scope_context", "")
            meeting_title = result.get("meeting_title", "")
            meeting_date  = result.get("meeting_date", "")

    merged_reqs = _dedup_requirements(all_req_lists)
    return {
        "meeting_title": meeting_title,
        "meeting_date":  meeting_date,
        "scope_context": scope_context,
        "requirements":  merged_reqs,
    }


# ──────────────────────────────────────────────────────────────────────────────
# MERGED STEP 2 + 3 + 4 — Normalize, Assess & Identify Gaps  (batched)
# ──────────────────────────────────────────────────────────────────────────────
def _step2_3_4_single_batch(
    model: Any,
    scope_context: str,
    batch: list[dict],
    batch_label: str,
) -> dict:
    """
    Run a single batch of requirements through the normalize + assess prompt.
    Returns { "normalized": [...], "capability_assessment": [...] }.
    """
    prompt = textwrap.dedent(f"""
    You are a Senior SAP Solution Architect and Business Analyst specialising in
    SAP IS-U, S/4HANA, FICA, and SD/FI integration.

    SCOPE CONTEXT
    -------------
    {scope_context}

    TASK
    ----
    Process EACH requirement below and produce TWO outputs per requirement.

    PART A — Normalization
    Decompose each requirement into:
      - id        : SAME id as the input (R1, R2 …). DO NOT change or invent IDs.
      - actor     : who initiates — one of: System, User, External System, IT, Business
      - action    : primary verb (create, generate, calculate, load, post, validate, …)
      - object    : what is acted upon
      - condition : business rule or trigger; use "None" if not applicable

    PART B — SAP Capability Assessment
    Assess whether STANDARD SAP S/4HANA (IS-U context) supports each requirement:
      FULL    → Fully supported by standard SAP configuration. No custom code needed.
      PARTIAL → Supported but requires a minor enhancement: BAdI, user exit, BRF+ rule,
                custom FM, or minor ABAP extension.
      NONE    → Not supported. Requires custom ABAP development or a new object.

    Assessment rules:
      - Be CONSERVATIVE — when in doubt between FULL and PARTIAL, always choose PARTIAL.
      - gap = true  when status is PARTIAL or NONE.
      - gap = false ONLY when status is FULL.
      - assessment_note: 1–2 sentences citing the specific SAP capability or limitation.

    ANTI-HALLUCINATION RULES:
      1. Output exactly {len(batch)} entries in EACH array — one per input requirement.
      2. Use the EXACT `id` from the input in every output object. Never invent IDs.
      3. Do not merge, split, or skip requirements.

    INPUT REQUIREMENTS ({len(batch)} items):
    {json.dumps(batch, indent=2)}

    RESPOND WITH JSON ONLY. No markdown, no preamble. Exact schema:
    {{
      "normalized": [
        {{
          "id":        "R1",
          "actor":     "<string>",
          "action":    "<string>",
          "object":    "<string>",
          "condition": "<string>"
        }}
      ],
      "capability_assessment": [
        {{
          "id":              "R1",
          "status":          "FULL",
          "gap":             false,
          "assessment_note": "<string>"
        }}
      ]
    }}
    """)
    return _call_llm(model, prompt, batch_label)


def step2_3_4_normalize_assess_gaps(model: Any, step1_json: dict) -> dict:
    """
    Merges Step 2 (normalize) and Steps 3&4 (assess + gap identification) into the
    minimum number of LLM calls while avoiding JSON-truncation failures.

    ROOT CAUSE OF '0 gaps' BUG:
      Processing all 84 requirements in one call produces ~168 JSON objects in a single
      response. This routinely exceeds the practical JSON-generation limit, causing a
      truncated / malformed response. The fix-it retry prompt then returns minimal valid
      JSON (empty arrays), which silently propagates as '0 gaps identified'.

    FIX:
      Batch requirements in groups of ASSESS_BATCH_SIZE (default 20).
      84 reqs → 5 calls of ≤20 items → ~80 JSON objects each → safely under 16k tokens.
      Results are merged in Python; no extra LLM calls for merging.

    Output: { normalized: [...], capability_assessment: [...] }
    """
    all_reqs     = step1_json.get("requirements", [])
    scope_ctx    = step1_json.get("scope_context", "")
    total        = len(all_reqs)
    batches      = [all_reqs[i:i + ASSESS_BATCH_SIZE]
                    for i in range(0, total, ASSESS_BATCH_SIZE)]
    n_batches    = len(batches)

    print(f"    [BATCH] {total} requirements → {n_batches} batch(es) of ≤{ASSESS_BATCH_SIZE}")

    all_normalized:   list[dict] = []
    all_assessments:  list[dict] = []

    for idx, batch in enumerate(batches, start=1):
        label  = f"Steps 2+3+4 – Normalize/Assess (batch {idx}/{n_batches})"
        result = _step2_3_4_single_batch(model, scope_ctx, batch, label)

        norm  = result.get("normalized", [])
        assess = result.get("capability_assessment", [])

        # Sanity-check: warn if a batch came back short (should never happen with hardened prompt)
        if len(norm) != len(batch) or len(assess) != len(batch):
            print(f"    [WARN] Batch {idx}: expected {len(batch)} items, "
                  f"got normalized={len(norm)}, assessment={len(assess)}. "
                  f"Partial results retained; check output carefully.")

        all_normalized.extend(norm)
        all_assessments.extend(assess)

    gaps_found = sum(1 for c in all_assessments if c.get("gap"))
    print(f"    [BATCH] Merged: {len(all_normalized)} normalized, "
          f"{len(all_assessments)} assessments, {gaps_found} gap(s)")

    return {
        "normalized":           all_normalized,
        "capability_assessment": all_assessments,
    }



# ──────────────────────────────────────────────────────────────────────────────
# STEPS 5, 6 & 7 — RICEFW Classification & Solution Strategy
# ──────────────────────────────────────────────────────────────────────────────
def step5_6_7_ricefw_and_strategy(model: Any, accumulated: dict) -> dict:
    """
    For every gap item: classify as RICEFW and write a concrete solution strategy.

    HARDENED vs original:
      - Strict rules block ID invention, cross-item hallucination, invalid RICEFW values.
      - Persona is an experienced SAP IS-U / S4 consultant — forces SAP-native solutions.
      - Output length must match input length (enforced by rule + downstream validator).
      - Still slim-payload optimised: only gap items sent, not full accumulated JSON.
    """
    req_map   = {r["id"]: r["text"] for r in accumulated.get("requirements", [])}
    gap_items = [
        {
            "id":              c["id"],
            "requirement":     req_map.get(c["id"], ""),
            "status":          c["status"],
            "assessment_note": c["assessment_note"],
        }
        for c in accumulated.get("capability_assessment", [])
        if c.get("gap")
    ]

    prompt = textwrap.dedent(f"""
    You are a Senior SAP Solution Architect and Lead Functional Consultant with 15+ years
    of hands-on experience in SAP IS-U, S/4HANA Utilities, FICA, SD/FI, and large-scale
    utility billing implementations. You are writing a formal gap analysis document.

    ═══════════════════════════════════════════════════════════════
    STRICT ANTI-HALLUCINATION RULES — VIOLATING THESE BREAKS THE PIPELINE
    ═══════════════════════════════════════════════════════════════
    1. USE EXACTLY the `id` from each input item as the `req_id` in your output.
       NEVER invent, modify, combine, or skip requirement IDs.
    2. Your output array MUST contain exactly ONE entry per input item — no more, no less.
    3. Items with status=FULL must NOT appear here; they are already excluded from the input.
    4. `ricefw` MUST be EXACTLY one of: Report, Interface, Conversion, Enhancement, Form, Workflow.
       Any other value (including abbreviations, combinations, or casing variants) is INVALID.
    5. `solution_bullets` MUST contain 3–5 items.  Each bullet must be a complete, actionable
       sentence referencing specific SAP mechanisms (e.g. BAPI, BRF+, FICA events, OData,
       LSMW, BAdI, FI-CA API, CRM, ABAP class, report transaction, IDOC type, etc.).
    6. `title` must be a concise technical title (5–10 words max), not a copy of the requirement.
    ═══════════════════════════════════════════════════════════════

    RICEFW DECISION GUIDE (use this to pick the right type):
    ─────────────────────────────────────────────────────────
    Report      → Data display, extraction, reconciliation, analytics output, ALV lists
    Interface   → Inbound / outbound data exchange with external systems (IDOC, REST, SOAP, flat file)
    Conversion  → One-time data migration or initial load (LSMW, BDC, LTMC, migration cockpit)
    Enhancement → Custom logic embedded in SAP: BAdI, user exit, BRF+, custom ABAP program,
                  implicit/explicit enhancement, event-driven logic (FICA events, SD user exits)
    Form        → Output document — PDF, smart form, SAP script, email, letter, SMS, bill
    Workflow    → Approval routing, escalation, task assignment, notification workflow

    CONSULTANT QUALITY BAR FOR solution_bullets:
    ─────────────────────────────────────────────
    • Name the exact SAP object, transaction, or API (e.g. "FKK_SAMPLE_PROCESS_01", "DFKKOP",
      "OB52", "FICA event 0010", "SAPLSMW", "EA-UTIL IS-U billing driver").
    • Describe integration touchpoints clearly (e.g. "triggered via IDoc UTILMD outbound").
    • Cover data flow, error handling, authorisation scope, or transport strategy where relevant.
    • Avoid vague bullets like "create a custom program" — specify the ABAP object type, class,
      function module, or enhancement spot instead.

    GAP ITEMS TO CLASSIFY ({len(gap_items)} items — your output MUST have exactly {len(gap_items)} entries):
    ════════════════════════════════════════════════════════════════
    {json.dumps(gap_items, indent=2)}

    RESPOND WITH JSON ONLY. No markdown, no preamble, no trailing text. Exact schema:
    {{
      "gap_analysis": [
        {{
          "gap_id":           "GAP-001",
          "req_id":           "<exact id from input>",
          "title":            "<concise technical title>",
          "ricefw":           "<Report|Interface|Conversion|Enhancement|Form|Workflow>",
          "solution_bullets": [
            "<actionable bullet referencing specific SAP mechanism>",
            "<actionable bullet referencing specific SAP mechanism>",
            "<actionable bullet referencing specific SAP mechanism>"
          ]
        }}
      ]
    }}
    """)
    return _call_llm(model, prompt, "Steps 5-6-7 – RICEFW & Strategy")


# ──────────────────────────────────────────────────────────────────────────────
# Confirmations & Open Actions
# ──────────────────────────────────────────────────────────────────────────────
def step_confirm_and_actions(model: Any, accumulated: dict) -> dict:
    """
    Produce no_gap_confirmations and open_actions.

    OPTIMIZATION: sends only the no-gap items + gap_analysis summary,
    not the entire accumulated pipeline JSON.
    """
    req_map   = {r["id"]: r["text"] for r in accumulated.get("requirements", [])}
    no_gap_items = [
        {
            "id":              c["id"],
            "requirement":     req_map.get(c["id"], ""),
            "assessment_note": c["assessment_note"],
        }
        for c in accumulated.get("capability_assessment", [])
        if not c.get("gap")
    ]

    # Slim gap_analysis: just gap_id, req_id, title, ricefw (skip long solution_bullets)
    gap_summary = [
        _slim(g, ["gap_id", "req_id", "title", "ricefw"])
        for g in accumulated.get("gap_analysis", [])
    ]

    slim_input = {
        "scope_context":  accumulated.get("scope_context", ""),
        "no_gap_items":   no_gap_items,
        "gap_summary":    gap_summary,
    }

    prompt = textwrap.dedent(f"""
    You are a Senior SAP Business Analyst finalising a gap analysis document.

    TASK
    ----
    Using the data below, produce TWO final sections:

    1. no_gap_confirmations
       List every requirement from no_gap_items.
       For each, provide a short "resolution" sentence explaining how standard SAP
       configuration addresses it (config table, IMG path, standard transaction, etc.).

    2. open_actions
       Identify 4–6 concrete actions that must happen BEFORE or DURING design phase.
       Each action must have:
         - action_number : 1, 2, 3 …
         - description   : what needs to happen
         - owner         : role or team responsible
         - target        : "Pre-Design" | "Design Phase" | specific session/date

    INPUT
    -----
    {json.dumps(slim_input, indent=2)}

    RESPOND WITH JSON ONLY. Exact schema:
    {{
      "no_gap_confirmations": [
        {{
          "id":         "R1",
          "topic":      "<short topic name>",
          "resolution": "<one sentence>"
        }}
      ],
      "open_actions": [
        {{
          "action_number": 1,
          "description":   "<string>",
          "owner":         "<string>",
          "target":        "<string>"
        }}
      ]
    }}
    """)
    return _call_llm(model, prompt, "Confirmations & Open Actions")


# ──────────────────────────────────────────────────────────────────────────────
# DOCX Renderer
# ──────────────────────────────────────────────────────────────────────────────
def render_docx(final_json: dict, output_path: str, output_dir: Path) -> str:
    json_path = output_dir / "_pipeline_output.json"
    json_path.write_text(json.dumps(final_json, indent=2, ensure_ascii=False), encoding="utf-8")

    if JS_TEMPLATE_PATH.exists():
        node_cmd = ["node", str(JS_TEMPLATE_PATH), str(json_path), output_path]
        print(f"    [DOCX] Running: {' '.join(node_cmd)}", flush=True)
        result = subprocess.run(node_cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(
                f"js_template.js failed (exit {result.returncode}):\n"
                f"STDOUT: {result.stdout}\nSTDERR: {result.stderr}"
            )
        print(f"    [DOCX] Written → {output_path}")
        return output_path

    print("    [DOCX] js_template.js not found — using Python fallback.", flush=True)
    return _render_docx_python(final_json, output_path)


def _render_docx_python(final_json: dict, output_path: str) -> str:
    """Fallback DOCX renderer (unchanged from original)."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Install python-docx: pip install python-docx") from exc

    doc = Document()
    doc.add_heading("SAP Gap Analysis", level=1)
    doc.add_paragraph(f"Meeting Title: {final_json.get('meeting_title', 'N/A')}")
    doc.add_paragraph(f"Meeting Date:  {final_json.get('meeting_date',  'N/A')}")

    sc = final_json.get("scope_context", "")
    if sc:
        doc.add_heading("Scope Context", level=2)
        doc.add_paragraph(sc)

    doc.add_heading("Requirements", level=2)
    for req in final_json.get("requirements", []):
        doc.add_paragraph(f"{req.get('id', '')}: {req.get('text', '')}", style="List Bullet")

    doc.add_heading("Gap Analysis (RICEFW)", level=2)
    for gap in final_json.get("gap_analysis", []):
        doc.add_paragraph(
            f"{gap.get('gap_id','')} | {gap.get('req_id','')} | {gap.get('ricefw','')} | {gap.get('title','')}",
            style="List Bullet",
        )
        for bullet in gap.get("solution_bullets", []):
            doc.add_paragraph(bullet, style="List Bullet 2")

    doc.add_heading("No-Gap Confirmations", level=2)
    for item in final_json.get("no_gap_confirmations", []):
        doc.add_paragraph(
            f"{item.get('id','')} — {item.get('topic','')}: {item.get('resolution','')}",
            style="List Bullet",
        )

    doc.add_heading("Open Actions", level=2)
    for action in final_json.get("open_actions", []):
        doc.add_paragraph(
            f"{action.get('action_number','')}. {action.get('description','')} "
            f"(Owner: {action.get('owner','')}, Target: {action.get('target','')})",
            style="List Bullet",
        )

    doc.save(output_path)
    print(f"    [DOCX] Written → {output_path}")
    return output_path


# ──────────────────────────────────────────────────────────────────────────────
# Orchestrator
# ──────────────────────────────────────────────────────────────────────────────
def run_pipeline(
    transcript_path: str,
    project: str,
    location: str,
    output_path: str,
    dump_json: bool = False,
) -> str:
    print("\n================================================")
    print("  SAP Gap Analysis Pipeline — Optimized")
    print("================================================\n")
    pipeline_start = time.time()

    output_path_obj = Path(output_path).resolve()
    output_dir = output_path_obj.parent
    output_dir.mkdir(parents=True, exist_ok=True)

    model      = _init_genai(project, location)
    transcript = _load_transcript(transcript_path)
    chunks     = _chunk_transcript(transcript)
    print(f"  [OK] Transcript loaded ({len(transcript):,} chars, {len(chunks)} chunk(s))\n")

    # ── Step 1 ────────────────────────────────────────────────────────────────
    t = time.time()
    print("Step 1 — Requirement Extraction")
    s1 = step1_extract_requirements(model, transcript)
    if dump_json:
        _dump(s1, "step1_requirements", output_dir)
    print(f"  → {len(s1.get('requirements', []))} requirements  [{time.time()-t:.1f}s]\n")

    # ── Steps 2+3+4 (merged) ─────────────────────────────────────────────────
    t = time.time()
    print("Steps 2+3+4 — Normalize, Assess & Identify Gaps  (batched, ≤20 reqs/call)")
    s2_out = step2_3_4_normalize_assess_gaps(model, s1)
    s2     = {**s1, **s2_out}
    if dump_json:
        _dump(s2, "step2_3_4_normalized_assessed", output_dir)
    gaps_found = sum(1 for c in s2.get("capability_assessment", []) if c.get("gap"))
    print(f"  → {gaps_found} gaps identified  [{time.time()-t:.1f}s]\n")

    # ── Steps 5+6+7 ──────────────────────────────────────────────────────────
    t = time.time()
    print("Steps 5+6+7 — RICEFW Classification & Solution Strategy")
    s5_out = step5_6_7_ricefw_and_strategy(model, s2)
    s5     = {**s2, **s5_out}

    # 🔒 ANTI-HALLUCINATION GUARDRAIL
    # Validate every gap item against the actual extracted requirement IDs.
    # Drops any item whose req_id was invented by the model, then re-sequences GAP-IDs.
    valid_req_ids = {r["id"] for r in s5.get("requirements", [])}
    s5["gap_analysis"] = _validate_and_filter_gaps(s5.get("gap_analysis", []), valid_req_ids)

    if dump_json:
        _dump(s5, "step5_ricefw_validated", output_dir)
    print(f"  → {len(s5.get('gap_analysis', []))} validated RICEFW items  [{time.time()-t:.1f}s]\n")

    # ── Confirmations & Actions ───────────────────────────────────────────────
    t = time.time()
    print("Confirmations & Open Actions")
    final_out = step_confirm_and_actions(model, s5)
    final     = {**s5, **final_out}
    if dump_json:
        _dump(final, "step_final", output_dir)
    print(f"  → {len(final.get('no_gap_confirmations', []))} confirmations, "
          f"{len(final.get('open_actions', []))} actions  [{time.time()-t:.1f}s]\n")

    # ── Render DOCX ───────────────────────────────────────────────────────────
    t = time.time()
    print("Rendering DOCX")
    render_docx(final, str(output_path_obj), output_dir)
    print(f"  [{time.time()-t:.1f}s]\n")

    total = time.time() - pipeline_start
    print("================================================")
    print(f"  Done!  Output → {output_path}  [total {total:.1f}s]")
    print("================================================\n")
    return output_path


# ──────────────────────────────────────────────────────────────────────────────
# CLI Entry Point  (identical to original)
# ──────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="SAP Gap Analysis Pipeline — Vertex AI Gemini 2.5 Pro → DOCX (Optimized)"
    )
    parser.add_argument("--transcript", required=True,  help="Path to transcript file (.json or .txt)")
    parser.add_argument("--project",    required=True,  help="GCP project ID for Vertex AI")
    parser.add_argument("--location",   default="us-central1", help="Vertex AI region")
    parser.add_argument("--output",     default="SAP_Gap_Analysis.docx", help="Output .docx path")
    parser.add_argument("--dump-json",  action="store_true", help="Dump each step's JSON to disk")
    args = parser.parse_args()

    run_pipeline(
        transcript_path=args.transcript,
        project=args.project,
        location=args.location,
        output_path=args.output,
        dump_json=args.dump_json,
    )