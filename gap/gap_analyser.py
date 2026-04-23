"""
SAP Gap Analysis Pipeline
=========================
Vertex AI · Gemini 2.5 Pro · 6 functions · JSON-chained steps → DOCX

Architecture
------------
  transcript (text/json)
      │
      ▼
  step1_extract_requirements()        → requirements_json
      │
      ▼
  step2_normalize_requirements()      → normalized_json        (receives requirements_json)
      │
      ▼
  step3_4_assess_and_identify_gaps()  → capability_json        (receives normalized_json)
      │
      ▼
  step5_6_7_ricefw_and_strategy()     → gap_analysis_json      (receives capability_json)
      │
      ▼
  step_confirm_and_actions()          → final_json             (receives gap_analysis_json)
      │
      ▼
  render_docx()                       → output.docx            (calls js_template.js via Node)

Usage
-----
  python sap_gap_pipeline.py \
      --transcript  meeting_input.json \
      --project     YOUR_GCP_PROJECT_ID \
      --location    us-central1 \
      --output      SAP_Gap_Analysis.docx

  # Optionally dump intermediate JSON:
      --dump-json

Environment
-----------
  GOOGLE_APPLICATION_CREDENTIALS=/path/to/service_account.json   (or ADC)
"""

import argparse
import json
import os
import subprocess
import sys
import textwrap
from datetime import datetime
from pathlib import Path

import vertexai
from vertexai.generative_models import GenerativeModel, GenerationConfig

# ──────────────────────────────────────────────────────────────────────────────
# Configuration
# ──────────────────────────────────────────────────────────────────────────────
MODEL_ID         = "gemini-2.5-pro"
TEMPERATURE      = 0.1      # near-deterministic for analysis tasks
MAX_TOKENS       = 12000
SCRIPT_DIR       = Path(__file__).parent
JS_TEMPLATE_PATH = SCRIPT_DIR / "js_template.js"


# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────
def _init_vertex(project: str, location: str) -> GenerativeModel:
    """Initialise Vertex AI and return a Gemini 2.5 Pro model instance."""
    vertexai.init(project=project, location=location)
    return GenerativeModel(
        MODEL_ID,
        generation_config=GenerationConfig(
            temperature=TEMPERATURE,
            max_output_tokens=MAX_TOKENS,
            response_mime_type="application/json",   # force JSON output
        ),
    )


def _call_llm(model: GenerativeModel, prompt: str, step_name: str) -> dict:
    """
    Send a prompt to Gemini 2.5 Pro and return parsed JSON.
    The model is instructed to respond ONLY with valid JSON.
    Raises ValueError if the response cannot be parsed.
    """
    print(f"  [LLM] Calling Gemini 2.5 Pro - {step_name} ...", flush=True)
    prompt_to_send = prompt
    last_raw = ""
    last_error = None

    for attempt in range(1, 4):
        response = model.generate_content(prompt_to_send)
        raw = response.text.strip()
        last_raw = raw

        # Strip optional markdown fences if the model ignores response_mime_type
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1]
            raw = raw.rsplit("```", 1)[0]

        try:
            return json.loads(raw)
        except json.JSONDecodeError as exc:
            last_error = exc
            if attempt < 3:
                print(f"  [LLM] Invalid JSON on attempt {attempt}; retrying...", flush=True)
                prompt_to_send = textwrap.dedent(f"""

                Your previous response was invalid JSON and could not be parsed.
                Return ONLY valid JSON with proper escaping and all strings terminated.
                Do not include markdown fences and do not add any explanation.

                ORIGINAL TASK:
                {prompt}
                
                """)
                continue

    raise ValueError(
        f"[{step_name}] LLM returned non-JSON output after retries.\n"
        f"Error: {last_error}\nRaw (first 500 chars): {last_raw[:500]}"
    ) from last_error


def _load_transcript(path: str) -> str:
    """Load and flatten a transcript file (JSON array or plain text)."""
    data = Path(path).read_text(encoding="utf-8")
    try:
        obj = json.loads(data)
        # EMS-style: array of meetings with transcript arrays.
        # Combine all meetings into one analysis corpus instead of only the first item.
        if isinstance(obj, list) and obj and isinstance(obj[0], dict) and "transcript_json" in obj[0]:
            chunks = []
            for idx, meeting in enumerate(obj, start=1):
                transcript_json = meeting.get("transcript_json", {})
                file_metadata = transcript_json.get("file_metadata", {})
                segments = transcript_json.get("transcript", [])
                lines = [s.get("text", "").strip() for s in segments if s.get("text", "").strip()]
                if not lines:
                    continue

                title = file_metadata.get("source_video") or file_metadata.get("file_name") or f"Meeting {idx}"
                meeting_number = meeting.get("meeting_number", idx)
                meeting_date = meeting.get("meeting_date", "unknown-date")
                chunks.append(
                    "\n".join([
                        f"=== Meeting {meeting_number} | {meeting_date} | {title} ===",
                        *lines,
                    ])
                )

            if chunks:
                return "\n\n".join(chunks)
        return json.dumps(obj, indent=2)
    except json.JSONDecodeError:
        return data          # plain text fallback


def _dump(data: dict, name: str) -> None:
    """Write a JSON step output to disk for inspection."""
    out = SCRIPT_DIR / f"{name}.json"
    out.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"  [DUMP] {out}")


# ──────────────────────────────────────────────────────────────────────────────
# STEP 1 — Requirement Extraction
# ──────────────────────────────────────────────────────────────────────────────
def step1_extract_requirements(model: GenerativeModel, transcript: str) -> dict:
    """
    Extract atomic business requirements from the raw meeting transcript.

    Input  : raw transcript text
    Output : { "requirements": [ { "id": "R1", "text": "..." }, ... ],
               "scope_context": "..." }
    """
    prompt = textwrap.dedent(f"""
    You are a Senior SAP Business Analyst performing requirement extraction from a meeting transcript.

    TASK
    ----
    Read the transcript carefully. Extract ALL distinct, actionable business requirements.

    RULES
    -----
    - Each requirement must be ATOMIC (one clear idea).
    - Use "The system should …" or "The business requires …" phrasing.
    - Ignore pleasantries, repetitions, off-topic chat.
    - Number requirements R1, R2, R3 … in order of appearance.
    - Also write a short "scope_context" paragraph (3–5 sentences) summarising what the
      session was about, the system landscape (IS-U / EMS / ECC), and key design decisions confirmed.

    RESPOND WITH JSON ONLY — no prose, no markdown fences. Schema:
    {{
      "meeting_title": "<string>",
      "meeting_date":  "<string>",
      "scope_context": "<string>",
      "requirements":  [
        {{ "id": "R1", "text": "<string>" }},
        ...
      ]
    }}

    TRANSCRIPT
    ----------
    {transcript[:30000]}
    """)
    return _call_llm(model, prompt, "Step 1 – Requirement Extraction")


# ──────────────────────────────────────────────────────────────────────────────
# STEP 2 — Requirement Normalization
# ──────────────────────────────────────────────────────────────────────────────
def step2_normalize_requirements(model: GenerativeModel, step1_json: dict) -> dict:
    """
    Normalize each requirement into Actor / Action / Object / Condition.

    Input  : step1_json (contains requirements list)
    Output : step1_json extended with "normalized" list
    """
    prompt = textwrap.dedent(f"""
    You are a Senior SAP Business Analyst performing requirement normalization.

    TASK
    ----
    For EACH requirement in the input JSON, produce a normalized record with:
      - id       : same as input (R1, R2 …)
      - actor    : who initiates (System, User, External System, IT, Business)
      - action   : verb (create, generate, calculate, load, post, …)
      - object   : what is acted upon
      - condition: business rule or trigger (if any); use "None" if not applicable

    Remove ambiguity. Be precise. Do NOT invent new requirements.

    INPUT JSON
    ----------
    {json.dumps(step1_json, indent=2)}

    RESPOND WITH JSON ONLY. Return ONLY this schema:
    {{
      "normalized": [
        {{
          "id":        "R1",
          "actor":     "<string>",
          "action":    "<string>",
          "object":    "<string>",
          "condition": "<string>"
        }},
        ...
      ]
    }}
    """)
    return _call_llm(model, prompt, "Step 2 – Normalization")


# ──────────────────────────────────────────────────────────────────────────────
# STEPS 3 & 4 — SAP Capability Assessment & Gap Identification
# ──────────────────────────────────────────────────────────────────────────────
def step3_4_assess_and_identify_gaps(model: GenerativeModel, step2_json: dict) -> dict:
    """
    Assess standard SAP S/4HANA IS-U capability for each requirement.
    Identify gaps (PARTIAL or NONE → gap=true).

    Input  : step2_json
    Output : step2_json extended with "capability_assessment" list
    """
    prompt = textwrap.dedent(f"""
    You are a Senior SAP Solution Architect specialising in SAP IS-U (Industry Solution Utilities),
    S/4HANA, FICA (Financial Contract Accounting), and SD/FI integration.

    TASK
    ----
    For EACH requirement, assess whether STANDARD SAP S/4HANA (IS-U context) supports it:
      - FULL    : Fully supported by standard SAP configuration/customising. No custom code.
      - PARTIAL : Supported but requires minor enhancement, BAdI, user exit, or BRF+ logic.
      - NONE    : Not supported. Custom ABAP development required.

    RULES
    -----
    - Be CONSERVATIVE. When in doubt, use PARTIAL.
    - Do NOT assume SAP behaviour without solid IS-U / FICA reasoning.
    - gap = true  when status is PARTIAL or NONE.
    - gap = false when status is FULL.
    - Provide a concise "assessment_note" (1–2 sentences) explaining your decision.

    INPUT JSON
    ----------
    {json.dumps(step2_json, indent=2)}

    RESPOND WITH JSON ONLY. Return ONLY this schema:
    {{
      "capability_assessment": [
        {{
          "id":              "R1",
          "status":          "FULL | PARTIAL | NONE",
          "gap":             true | false,
          "assessment_note": "<string>"
        }},
        ...
      ]
    }}
    """)
    return _call_llm(model, prompt, "Steps 3&4 – Capability & Gap")


# ──────────────────────────────────────────────────────────────────────────────
# STEPS 5, 6 & 7 — RICEFW Classification, Solution Strategy, Final GAP Table
# ──────────────────────────────────────────────────────────────────────────────
def step5_6_7_ricefw_and_strategy(model: GenerativeModel, step3_json: dict) -> dict:
    """
    For every gap item: classify as RICEFW and write a concrete solution strategy.
    Produce the final GAP analysis table.

    Input  : step3_json
    Output : step3_json extended with "gap_analysis" list
    """
    prompt = textwrap.dedent(f"""
    You are a Senior SAP Solution Architect and Functional Consultant.

    TASK
    ----
    Process ONLY requirements where gap=true in the capability_assessment.

    For each gap:
    1. RICEFW Classification — assign exactly ONE type:
         Report      : Data display, extraction, analytics, reconciliation output
         Interface   : Inbound or outbound data exchange between systems
         Conversion  : One-time data migration or initial load
         Enhancement : Custom logic, BAdI, user exit, BRF+, custom program
         Form        : Output document — PDF, email, letter, SMS
         Workflow    : Approval routing, task assignment

    2. Solution Strategy — implementation-oriented description covering:
         - What will be built (3–5 bullet points)
         - Trigger / event (batch job, FICA event, report run, etc.)
         - SAP mechanisms used (BRF+, FICA API, LSMW, BAPI, etc.)
         - Integration touchpoints

    RULES
    -----
    - Use strict RICEFW logic. Do NOT guess.
    - Each bullet in solution_bullets must be a complete, actionable sentence.
    - gap_id format: GAP-001, GAP-002, … (sequential, not matching R-numbers).

    INPUT JSON
    ----------
    {json.dumps(step3_json, indent=2)}

    RESPOND WITH JSON ONLY. Return ONLY this schema:
    {{
      "gap_analysis": [
        {{
          "gap_id":           "GAP-001",
          "req_id":           "R2",
          "title":            "<short meaningful title>",
          "ricefw":           "Report | Interface | Conversion | Enhancement | Form | Workflow",
          "solution_bullets": ["<bullet 1>", "<bullet 2>", ...]
        }},
        ...
      ]
    }}
    """)
    return _call_llm(model, prompt, "Steps 5-6-7 – RICEFW & Strategy")


# ──────────────────────────────────────────────────────────────────────────────
# Confirmations & Open Actions
# ──────────────────────────────────────────────────────────────────────────────
def step_confirm_and_actions(model: GenerativeModel, step5_json: dict) -> dict:
    """
    Produce two final sections:
      - no_gap_confirmations : requirements covered by standard SAP config
      - open_actions         : actionable next steps with owner and target date

    Input  : step5_json (full pipeline JSON)
    Output : step5_json extended with "no_gap_confirmations" and "open_actions"
    """
    prompt = textwrap.dedent(f"""
    You are a Senior SAP Business Analyst finalising a gap analysis document.

    TASK
    ----
    Using the full pipeline JSON below, produce TWO final sections:

    1. no_gap_confirmations
       List every requirement where gap=false.
       For each, provide a short "resolution" sentence explaining how standard SAP
       configuration addresses it (config table, IMG path, standard transaction, etc.).

    2. open_actions
       Identify 4–6 concrete actions that must happen BEFORE or DURING design phase.
       Each action must have:
         - action_number : 1, 2, 3 …
         - description   : what needs to happen
         - owner         : role or team responsible
         - target        : "Pre-Design" | "Design Phase" | specific session/date

    INPUT JSON
    ----------
    {json.dumps(step5_json, indent=2)}

    RESPOND WITH JSON ONLY. Return ONLY this schema:
    {{
      "no_gap_confirmations": [
        {{
          "id":         "R1",
          "topic":      "<short topic name>",
          "resolution": "<one sentence>"
        }},
        ...
      ],
      "open_actions": [
        {{
          "action_number": 1,
          "description":   "<string>",
          "owner":         "<string>",
          "target":        "<string>"
        }},
        ...
      ]
    }}
    """)
    return _call_llm(model, prompt, "Confirmations & Open Actions")


# ──────────────────────────────────────────────────────────────────────────────
# DOCX Renderer  (Python → writes JSON → Node runs js_template.js → .docx)
# ──────────────────────────────────────────────────────────────────────────────
def render_docx(final_json: dict, output_path: str) -> str:
    """
    Serialise final_json to a temp file, then call:
        node js_template.js  <json_path>  <output_docx_path>
    Returns the resolved output path.
    """
    json_path = SCRIPT_DIR / "_pipeline_output.json"
    json_path.write_text(
        json.dumps(final_json, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )

    if JS_TEMPLATE_PATH.exists():
        node_cmd = ["node", str(JS_TEMPLATE_PATH), str(json_path), output_path]
        print(f"  [DOCX] Running: {' '.join(node_cmd)}", flush=True)

        result = subprocess.run(node_cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(
                f"js_template.js failed (exit {result.returncode}):\n"
                f"STDOUT: {result.stdout}\nSTDERR: {result.stderr}"
            )
        print(f"  [DOCX] Written -> {output_path}")
        return output_path

    print("  [DOCX] js_template.js not found. Using Python fallback renderer.", flush=True)
    return _render_docx_python(final_json, output_path)


def _render_docx_python(final_json: dict, output_path: str) -> str:
    """Fallback DOCX renderer when js_template.js is unavailable."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for fallback rendering. Install with: pip install python-docx"
        ) from exc

    doc = Document()
    doc.add_heading("SAP Gap Analysis", level=1)

    meeting_title = final_json.get("meeting_title", "N/A")
    meeting_date = final_json.get("meeting_date", "N/A")
    scope_context = final_json.get("scope_context", "")

    doc.add_paragraph(f"Meeting Title: {meeting_title}")
    doc.add_paragraph(f"Meeting Date: {meeting_date}")
    if scope_context:
        doc.add_heading("Scope Context", level=2)
        doc.add_paragraph(scope_context)

    doc.add_heading("Requirements", level=2)
    for req in final_json.get("requirements", []):
        doc.add_paragraph(f"{req.get('id', '')}: {req.get('text', '')}", style="List Bullet")

    doc.add_heading("Gap Analysis (RICEFW)", level=2)
    for gap in final_json.get("gap_analysis", []):
        doc.add_paragraph(
            f"{gap.get('gap_id', '')} | {gap.get('req_id', '')} | {gap.get('ricefw', '')} | {gap.get('title', '')}",
            style="List Bullet",
        )
        for bullet in gap.get("solution_bullets", []):
            doc.add_paragraph(bullet, style="List Bullet 2")

    doc.add_heading("No-Gap Confirmations", level=2)
    for item in final_json.get("no_gap_confirmations", []):
        topic = item.get("topic", "")
        resolution = item.get("resolution", "")
        doc.add_paragraph(f"{item.get('id', '')} - {topic}: {resolution}", style="List Bullet")

    doc.add_heading("Open Actions", level=2)
    for action in final_json.get("open_actions", []):
        doc.add_paragraph(
            f"{action.get('action_number', '')}. {action.get('description', '')} "
            f"(Owner: {action.get('owner', '')}, Target: {action.get('target', '')})",
            style="List Bullet",
        )

    doc.save(output_path)
    print(f"  [DOCX] Written -> {output_path}")
    return output_path


# ──────────────────────────────────────────────────────────────────────────────
# Orchestrator
# ──────────────────────────────────────────────────────────────────────────────
def run_pipeline(
    transcript_path: str,
    project: str,
    location: str,
    output_path: str,
    dump_json: bool = False,
) -> str:
    """
    Full pipeline:
      transcript → S1 → S2 → S3/4 → S5/6/7 → Confirm → DOCX
    Each step receives the previous step's JSON and enriches it.
    """
    print("\n==============================================")
    print("  SAP Gap Analysis Pipeline - Gemini 2.5 Pro")
    print("==============================================\n")

    # Init
    model = _init_vertex(project, location)
    transcript = _load_transcript(transcript_path)
    print(f"  [OK] Transcript loaded ({len(transcript):,} chars)\n")

    # Step 1
    print("Step 1 - Requirement Extraction")
    s1 = step1_extract_requirements(model, transcript)
    if dump_json:
        _dump(s1, "step1_requirements")
    print(f"  -> {len(s1.get('requirements', []))} requirements extracted\n")

    # Step 2
    print("Step 2 - Requirement Normalization")
    s2_out = step2_normalize_requirements(model, s1)
    s2 = {**s1, **s2_out}
    if dump_json:
        _dump(s2, "step2_normalized")
    print(f"  -> {len(s2.get('normalized', []))} normalized records\n")

    # Steps 3 & 4
    print("Steps 3 & 4 - SAP Capability Assessment & Gap Identification")
    s3_out = step3_4_assess_and_identify_gaps(model, s2)
    s3 = {**s2, **s3_out}
    if dump_json:
        _dump(s3, "step3_capability")
    gaps_found = sum(1 for c in s3.get("capability_assessment", []) if c.get("gap"))
    print(f"  -> {gaps_found} gaps identified\n")

    # Steps 5, 6, 7
    print("Steps 5, 6 & 7 - RICEFW Classification & Solution Strategy")
    s5_out = step5_6_7_ricefw_and_strategy(model, s3)
    s5 = {**s3, **s5_out}
    if dump_json:
        _dump(s5, "step5_ricefw")
    print(f"  -> {len(s5.get('gap_analysis', []))} RICEFW items\n")

    # Confirmations & Actions
    print("Confirmations & Open Actions")
    final_out = step_confirm_and_actions(model, s5)
    final = {**s5, **final_out}
    if dump_json:
        _dump(final, "step_final")
    print(f"  -> {len(final.get('no_gap_confirmations', []))} no-gap confirmations")
    print(f"  -> {len(final.get('open_actions', []))} open actions\n")

    # Render DOCX
    print("Rendering DOCX via js_template.js")
    render_docx(final, output_path)

    print("\n==============================================")
    print(f"  Done! Output -> {output_path}")
    print("==============================================\n")
    return output_path


# ──────────────────────────────────────────────────────────────────────────────
# CLI Entry Point
# ──────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="SAP Gap Analysis Pipeline — Vertex AI Gemini 2.5 Pro → DOCX"
    )
    parser.add_argument(
        "--transcript", required=True,
        help="Path to transcript file (.json or .txt)"
    )
    parser.add_argument(
        "--project", required=True,
        help="GCP project ID for Vertex AI"
    )
    parser.add_argument(
        "--location", default="us-central1",
        help="Vertex AI region (default: us-central1)"
    )
    parser.add_argument(
        "--output", default="SAP_Gap_Analysis.docx",
        help="Output .docx file path"
    )
    parser.add_argument(
        "--dump-json", action="store_true",
        help="Dump each step's JSON to disk for inspection"
    )
    args = parser.parse_args()

    run_pipeline(
        transcript_path=args.transcript,
        project=args.project,
        location=args.location,
        output_path=args.output,
        dump_json=args.dump_json,
    )